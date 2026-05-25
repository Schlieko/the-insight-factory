import os
import time
import json
from docx import Document
from docx.shared import Pt, RGBColor
from google import genai
from dotenv import load_dotenv

# --- 1. CONFIGURATION ---
target_dir = os.path.dirname(os.path.abspath(__file__))

load_dotenv(os.path.join(target_dir, ".env"))
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    print("❌ ERROR: GEMINI_API_KEY not found in .env file.")
    exit()

client = genai.Client(api_key=api_key)

# Using the powerful 3.1 Pro Preview for both planning and heavy writing!
planner_model = 'gemini-3.1-pro-preview'
writer_model = 'gemini-3.1-pro-preview'

def load_text_file(filename):
    filepath = os.path.join(target_dir, filename)
    try:
        with open(filepath, 'r', encoding='utf-8') as file:
            return file.read().strip()
    except FileNotFoundError:
        return None

def main():
    print("📂 Loading Master Document and Audit Report...")
    master_document = load_text_file("master_output.md")
    audit_report = load_text_file("audit_report.md")
    
    if not master_document or not audit_report:
        print("❌ Error: Ensure both Phase 1 and Phase 2 scripts have successfully run.")
        return

    # --- STEP 1: GENERATE THE MANUAL FALLBACK PROMPT ---
    print("\n📝 Step 1: Generating the Manual Web UI Fallback Prompt...")
    meta_prompt = f"""
    You are an expert Prompt Engineer. Based on the provided master document and conceptual audit, 
    write a highly optimized, comprehensive prompt that the user can copy and paste into an 
    autonomous Deep Research AI. 
    
    The prompt MUST instruct the AI to:
    1. Explore the broad strategic advantages and societal transformations of the core idea.
    2. Research potential solutions or historical precedents for the flaws identified in the audit.
    3. Write a massive, highly detailed report (approx 20 pages) structured with clear headings.
    4. Ensure the final report is written at an 8th-grade reading level.
    
    Output ONLY the text of the prompt. Do not include conversational filler.
    """
    try:
        prompt_response = client.models.generate_content(
            model=planner_model,
            contents=meta_prompt
        )
        fallback_prompt = prompt_response.text.strip()
        
        fallback_filepath = os.path.join(target_dir, "Manual_Deep_Research_Prompt.txt")
        with open(fallback_filepath, 'w', encoding='utf-8') as f:
            f.write("--- COPY AND PASTE THIS INTO GEMINI ADVANCED (DEEP RESEARCH) ---\n")
            f.write("--- DON'T FORGET TO ATTACH YOUR MASTER.MD AND AUDIT.MD FILES ---\n\n")
            f.write(fallback_prompt)
        print("✅ Saved fallback prompt to 'Manual_Deep_Research_Prompt.txt'")
    except Exception as e:
        print(f"⚠️ Error generating fallback prompt: {e}")

    # --- STEP 2: GENERATE THE OUTLINE (The Planner) ---
    print("\n🧠 Step 2: Generating Comprehensive Report Outline...")
    outline_prompt = f"""
    You are an expert Research Director. Based on the provided master document and conceptual audit, 
    create a highly detailed, 8-section outline for a comprehensive Deep Research Report.
    
    The report must explore the strategic advantages, societal transformations, and solutions to the flaws identified in the audit.
    
    OUTPUT FORMAT:
    You must output ONLY valid JSON in the following format:
    [
      {{"section_title": "1. Executive Summary", "focus_instructions": "Summarize the core concept..."}},
      {{"section_title": "2. Strategic Advantages", "focus_instructions": "Detail the competitive edge..."}}
    ]

    [START MASTER DOCUMENT]
    {master_document}
    [END MASTER DOCUMENT]
    
    [START AUDIT REPORT]
    {audit_report}
    [END AUDIT REPORT]
    """
    try:
        outline_response = client.models.generate_content(
            model=planner_model,
            contents=outline_prompt
        )
        raw_json = outline_response.text.strip().replace("```json", "").replace("```", "")
        outline = json.loads(raw_json)
        print(f"✅ Successfully generated an {len(outline)}-section outline.")
    except Exception as e:
        print(f"❌ Error generating outline: {e}")
        return

    # --- STEP 3: PREPARE THE WORD DOCUMENT ---
    print("\n🎨 Applying professional styling theme...")
    doc = Document()
    
    # Style the normal paragraph text
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(11)
    
    # Style the Title
    title_style = doc.styles['Title']
    title_style.font.name = 'Calibri'
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(17, 85, 204) # Executive Blue
    
    # Style the Section Headings
    h1_style = doc.styles['Heading 1']
    h1_style.font.name = 'Calibri'
    h1_style.font.size = Pt(16)
    h1_style.font.bold = True
    h1_style.font.color.rgb = RGBColor(17, 85, 204)
    
    doc.add_heading('Deep Research Report', 0)

    # --- STEP 4: GENERATE SECTIONS (The Writer Loop) ---
    print("\n✍️ Step 4: Generating Report Section by Section (This will take a few minutes)...")
    
    for index, section in enumerate(outline):
        title = section.get("section_title", f"Section {index + 1}")
        instructions = section.get("focus_instructions", "")
        
        print(f"   -> Drafting: {title}...")
        
        section_prompt = f"""
        You are an expert Technical Writer. Write a highly detailed, comprehensive section for a Deep Research Report.
        Write at an 8th-grade reading level. Expand deeply on the concepts. 
        Aim for 1,000 to 1,500 words for this section alone.
        
        CURRENT SECTION TITLE: {title}
        YOUR INSTRUCTIONS FOR THIS SECTION: {instructions}
        
        CRITICAL FORMATTING INSTRUCTIONS:
        1. DO NOT use raw Markdown links like [text](url). 
        2. DO NOT output long, ugly URLs in the text. If you must reference a source, weave the name of the publication naturally into the sentence.
        3. DO NOT use asterisks (**) for bolding. Write in clean, professional plain text.
        
        Use the context below to inform your writing.
        [CONTEXT]
        {master_document}
        {audit_report}
        [/CONTEXT]
        
        Output ONLY the text for this specific section. Do not include introductory filler.
        """

        try:
            section_response = client.models.generate_content(
                model=writer_model,
                contents=section_prompt
            )
            
            doc.add_heading(title, level=1)
            doc.add_paragraph(section_response.text.strip())
            
            time.sleep(5) 
            
        except Exception as e:
            print(f"⚠️ Error drafting '{title}': {e}")
            doc.add_heading(title, level=1)
            doc.add_paragraph(f"[ERROR GENERATING SECTION: {e}]")

    # --- STEP 5: SAVE THE MASTERPIECE ---
    output_filename = "Deep_Research_Report.docx"
    output_filepath = os.path.join(target_dir, output_filename)
    doc.save(output_filepath)
    print(f"\n🎉 Success! Your massive report has been saved as '{output_filename}'")

if __name__ == "__main__":
    main()